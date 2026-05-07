import csv
from datetime import datetime
import json
import os
import shutil
import textwrap
from collections import Counter, defaultdict

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from PIL import Image, ImageDraw, ImageFont


RESULT_DIRS = (
    "predictions",
    "metrics",
    "failure_cases",
    "visualizations",
    os.path.join("visualizations", "charts"),
    "reports",
)

REASONING_LABELS = (
    "compositional",
    "causal",
    "emotional",
    "temporal",
    "social",
    "spatial",
)


def _ensure_dir(path):
    os.makedirs(path, exist_ok=True)
    return path


def _slugify(value):
    safe = []
    for char in str(value):
        if char.isalnum():
            safe.append(char.lower())
        elif char in {"-", "_"}:
            safe.append(char)
        else:
            safe.append("_")
    return "".join(safe).strip("_") or "item"


def _normalize_text(value):
    text = " ".join(str(value or "").strip().lower().split())
    for char in ",.!?;:\"'()[]{}":
        text = text.replace(char, " ")
    return " ".join(text.split())


def _token_overlap(left, right):
    left_tokens = set(_normalize_text(left).split())
    right_tokens = set(_normalize_text(right).split())
    if not left_tokens or not right_tokens:
        return 0.0
    return len(left_tokens & right_tokens) / max(len(right_tokens), 1)


def _load_font(size):
    candidates = [
        "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeui.ttf",
    ]
    for candidate in candidates:
        if os.path.exists(candidate):
            try:
                return ImageFont.truetype(candidate, size=size)
            except OSError:
                continue
    return ImageFont.load_default()


class AdversarialArtifactWriter:
    def __init__(self, repo_root, results_root=None):
        self.repo_root = repo_root
        self.results_root = results_root or os.path.join(repo_root, "results")
        self.paths = {name: _ensure_dir(os.path.join(self.results_root, name)) for name in RESULT_DIRS}
        self.predictions_path = os.path.join(self.paths["predictions"], "predictions.json")
        self.metrics_json_path = os.path.join(self.paths["metrics"], "metrics.json")
        self.metrics_csv_path = os.path.join(self.paths["metrics"], "metrics.csv")
        self.report_path = os.path.join(self.paths["reports"], "MemEIC_Adversarial_Failure_Analysis.docx")

    def finalize(self, new_records):
        all_records = self._merge_predictions(new_records)
        metrics = self._compute_metrics(all_records)
        self._write_metrics(metrics)
        failure_assets = self._write_failure_assets(all_records)
        chart_paths = self._write_charts(metrics)
        self._write_report(all_records, metrics, failure_assets, chart_paths)
        return {
            "overall_accuracy": metrics["overall_accuracy"],
            "adversarial_failure_rate": metrics["adversarial_failure_rate"],
            "hallucination_rate": metrics["hallucination_rate"],
            "prediction_count": len(all_records),
        }

    def build_prediction_record(self, *, image, question, ground_truth, model_prediction, correct, reasoning_type, model_name, gap_num, source_record):
        reasoning_type = self._map_reasoning_type(reasoning_type)
        failure_type = self._infer_failure_type(
            correct=correct,
            reasoning_type=reasoning_type,
            prediction=model_prediction,
            ground_truth=ground_truth,
            source_record=source_record,
        )
        image_rel = os.path.relpath(image, self.repo_root).replace("\\", "/") if os.path.isabs(image) else str(image).replace("\\", "/")
        return {
            "image": image_rel,
            "question": question,
            "ground_truth": ground_truth,
            "model_prediction": model_prediction,
            "correct": bool(correct),
            "correctness": bool(correct),
            "failure_type": failure_type,
            "reasoning_category": "multi_hop",
            "reasoning_type": reasoning_type,
            "model_name": model_name,
            "gap_num": int(gap_num),
            "source_alt": source_record.get("alt", ""),
        }

    def _merge_predictions(self, new_records):
        existing = []
        if os.path.exists(self.predictions_path):
            with open(self.predictions_path, "r", encoding="utf-8") as handle:
                existing = json.load(handle)

        merged = {}
        for record in existing + list(new_records):
            key = (
                record.get("model_name", ""),
                record.get("gap_num", 0),
                record.get("image", ""),
                record.get("question", ""),
                record.get("reasoning_type", ""),
            )
            merged[key] = record

        all_records = list(merged.values())
        all_records.sort(key=lambda item: (item.get("model_name", ""), item.get("gap_num", 0), item.get("image", ""), item.get("question", "")))
        with open(self.predictions_path, "w", encoding="utf-8") as handle:
            json.dump(all_records, handle, indent=2, ensure_ascii=False)
        return all_records

    def _compute_metrics(self, records):
        total = len(records)
        correct = sum(1 for record in records if record["correct"])
        hallucinations = sum(1 for record in records if record["failure_type"] == "hallucination")
        failure_counts = Counter(record["failure_type"] for record in records if not record["correct"])
        by_type = defaultdict(list)
        by_model = defaultdict(list)
        by_gap = defaultdict(list)

        for record in records:
            by_type[record["reasoning_type"]].append(record)
            by_model[record["model_name"]].append(record)
            by_gap[str(record.get("gap_num", 0))].append(record)

        return {
            "prediction_count": total,
            "overall_accuracy": self._safe_rate(correct, total),
            "hallucination_rate": self._safe_rate(hallucinations, total),
            "adversarial_failure_rate": self._safe_rate(total - correct, total),
            "multi_hop_reasoning_accuracy": self._safe_rate(correct, total),
            "compositional_reasoning_accuracy": self._bucket_accuracy(by_type.get("compositional", [])),
            "causal_reasoning_accuracy": self._bucket_accuracy(by_type.get("causal", [])),
            "emotional_reasoning_accuracy": self._bucket_accuracy(by_type.get("emotional", [])),
            "temporal_reasoning_accuracy": self._bucket_accuracy(by_type.get("temporal", [])),
            "social_reasoning_accuracy": self._bucket_accuracy(by_type.get("social", [])),
            "spatial_reasoning_accuracy": self._bucket_accuracy(by_type.get("spatial", [])),
            "failure_distribution": dict(failure_counts),
            "reasoning_type_counts": {label: len(by_type.get(label, [])) for label in REASONING_LABELS},
            "category_accuracy": {label: self._bucket_accuracy(by_type.get(label, [])) for label in REASONING_LABELS},
            "model_performance": {model: self._bucket_accuracy(model_records) for model, model_records in by_model.items()},
            "gap_performance": {gap: self._bucket_accuracy(gap_records) for gap, gap_records in by_gap.items()},
            "dataset_statistics": {
                "unique_images": len({record["image"] for record in records}),
                "evaluated_questions": total,
                "models": sorted(by_model.keys()),
            },
        }

    def _write_metrics(self, metrics):
        with open(self.metrics_json_path, "w", encoding="utf-8") as handle:
            json.dump(metrics, handle, indent=2, ensure_ascii=False)

        rows = [
            ("prediction_count", metrics["prediction_count"]),
            ("overall_accuracy", metrics["overall_accuracy"]),
            ("hallucination_rate", metrics["hallucination_rate"]),
            ("adversarial_failure_rate", metrics["adversarial_failure_rate"]),
            ("multi_hop_reasoning_accuracy", metrics["multi_hop_reasoning_accuracy"]),
            ("compositional_reasoning_accuracy", metrics["compositional_reasoning_accuracy"]),
            ("causal_reasoning_accuracy", metrics["causal_reasoning_accuracy"]),
            ("emotional_reasoning_accuracy", metrics["emotional_reasoning_accuracy"]),
            ("temporal_reasoning_accuracy", metrics["temporal_reasoning_accuracy"]),
            ("social_reasoning_accuracy", metrics["social_reasoning_accuracy"]),
            ("spatial_reasoning_accuracy", metrics["spatial_reasoning_accuracy"]),
        ]
        with open(self.metrics_csv_path, "w", encoding="utf-8", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(["metric", "value"])
            writer.writerows(rows)

    def _write_failure_assets(self, records):
        failures = [record for record in records if not record["correct"]]
        failure_assets = []
        for index, record in enumerate(failures[:24], start=1):
            slug = f"{index:03d}_{_slugify(record['model_name'])}_{_slugify(record['reasoning_type'])}"
            original_path = os.path.join(self.repo_root, record["image"])
            copied_image = os.path.join(self.paths["failure_cases"], f"{slug}_original{os.path.splitext(original_path)[1] or '.png'}")
            overlay_path = os.path.join(self.paths["failure_cases"], f"{slug}_failure.png")
            panel_path = os.path.join(self.paths["visualizations"], f"{slug}_panel.png")
            if os.path.exists(original_path):
                shutil.copy2(original_path, copied_image)
                self._render_failure_panel(original_path, record, overlay_path, academic=False)
                self._render_failure_panel(original_path, record, panel_path, academic=True)
                failure_assets.append({
                    "original": copied_image,
                    "failure_panel": overlay_path,
                    "academic_panel": panel_path,
                    "record": record,
                })
        return failure_assets

    def _render_failure_panel(self, image_path, record, output_path, academic):
        image = Image.open(image_path).convert("RGB")
        canvas = Image.new("RGB", (1600, 980), color=(249, 247, 241) if academic else (255, 255, 255))
        draw = ImageDraw.Draw(canvas)

        title_font = _load_font(34 if academic else 30)
        subtitle_font = _load_font(24)
        body_font = _load_font(22)

        image.thumbnail((760, 760))
        canvas.paste(image, (60, 120))

        accent = (130, 35, 20)
        secondary = (55, 55, 55)
        header = "MemEIC Adversarial Failure Analysis" if academic else "Failure Case"
        draw.text((60, 40), header, fill=(25, 25, 25), font=title_font)
        draw.text((860, 120), f"Reasoning Type: {record['reasoning_type'].title()}", fill=accent, font=subtitle_font)
        draw.text((860, 160), f"Failure Type: {record['failure_type'].title()}", fill=accent, font=subtitle_font)

        blocks = [
            ("Question", record["question"]),
            ("Ground Truth", record["ground_truth"]),
            ("Model Prediction", record["model_prediction"] or "<empty>"),
            ("Why It Failed", self._failure_summary(record)),
        ]

        y = 220
        for label, text in blocks:
            draw.text((860, y), label, fill=(15, 15, 15), font=subtitle_font)
            y += 38
            for line in textwrap.wrap(str(text), width=42):
                draw.text((860, y), line, fill=secondary, font=body_font)
                y += 30
            y += 28

        canvas.save(output_path)

    def _write_charts(self, metrics):
        chart_dir = self.paths[os.path.join("visualizations", "charts")]
        chart_paths = {}
        plt.style.use("seaborn-v0_8-whitegrid")

        category_items = metrics["category_accuracy"]
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.bar(list(category_items.keys()), list(category_items.values()), color="#3f6c51")
        ax.set_ylim(0, 1)
        ax.set_ylabel("Accuracy")
        ax.set_title("Category-wise Accuracy")
        fig.tight_layout()
        chart_paths["category_accuracy"] = os.path.join(chart_dir, "category_accuracy.png")
        fig.savefig(chart_paths["category_accuracy"], dpi=200)
        plt.close(fig)

        failure_items = metrics["failure_distribution"] or {"no_failures": 0}
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.bar(list(failure_items.keys()), list(failure_items.values()), color="#b75438")
        ax.set_ylabel("Count")
        ax.set_title("Failure Distribution")
        ax.tick_params(axis="x", rotation=25)
        fig.tight_layout()
        chart_paths["failure_distribution"] = os.path.join(chart_dir, "failure_distribution.png")
        fig.savefig(chart_paths["failure_distribution"], dpi=200)
        plt.close(fig)

        reasoning_items = metrics["category_accuracy"]
        fig, ax = plt.subplots(figsize=(10, 6))
        failure_rates = [1 - value for value in reasoning_items.values()]
        positions = range(len(reasoning_items))
        ax.bar([position - 0.2 for position in positions], list(reasoning_items.values()), width=0.4, label="Accuracy", color="#375a7f")
        ax.bar([position + 0.2 for position in positions], failure_rates, width=0.4, label="Failure Rate", color="#d16d5b")
        ax.set_xticks(list(positions), list(reasoning_items.keys()))
        ax.set_ylim(0, 1)
        ax.set_title("Reasoning-type Comparison")
        ax.legend()
        fig.tight_layout()
        chart_paths["reasoning_comparison"] = os.path.join(chart_dir, "reasoning_type_comparison.png")
        fig.savefig(chart_paths["reasoning_comparison"], dpi=200)
        plt.close(fig)

        model_items = metrics["model_performance"] or {"current_model": metrics["overall_accuracy"]}
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.bar(list(model_items.keys()), list(model_items.values()), color="#6b8e23")
        ax.set_ylim(0, 1)
        ax.set_ylabel("Accuracy")
        ax.set_title("Model Performance Comparison")
        ax.tick_params(axis="x", rotation=20)
        fig.tight_layout()
        chart_paths["model_comparison"] = os.path.join(chart_dir, "model_performance_comparison.png")
        fig.savefig(chart_paths["model_comparison"], dpi=200)
        plt.close(fig)

        return chart_paths

    def _write_report(self, records, metrics, failure_assets, chart_paths):
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError as exc:
            raise ImportError("python-docx is required to generate the DOCX analysis report.") from exc

        document = Document()
        document.add_heading("MemEIC Adversarial Failure Analysis", 0)

        document.add_heading("Introduction", level=1)
        document.add_paragraph(
            "This report summarizes GPU-accelerated adversarial multi-hop reasoning evaluation results for MemEIC and baseline-compatible evaluation paths in the current repository."
        )

        document.add_heading("Benchmark Overview", level=1)
        document.add_paragraph(
            f"Evaluated {metrics['prediction_count']} reasoning prompts across {metrics['dataset_statistics']['unique_images']} images."
        )

        document.add_heading("Dataset Statistics", level=1)
        stats_table = document.add_table(rows=1, cols=2)
        stats_table.rows[0].cells[0].text = "Statistic"
        stats_table.rows[0].cells[1].text = "Value"
        for key, value in metrics["dataset_statistics"].items():
            row = stats_table.add_row().cells
            row[0].text = str(key)
            row[1].text = ", ".join(value) if isinstance(value, list) else str(value)

        document.add_heading("Evaluation Setup", level=1)
        document.add_paragraph("Inference used the repository's CUDA-enabled evaluation pipeline and reused the existing sequential edit workflow.")

        document.add_heading("Quantitative Metrics", level=1)
        metric_table = document.add_table(rows=1, cols=2)
        metric_table.rows[0].cells[0].text = "Metric"
        metric_table.rows[0].cells[1].text = "Value"
        metric_names = [
            "overall_accuracy",
            "hallucination_rate",
            "adversarial_failure_rate",
            "multi_hop_reasoning_accuracy",
            "compositional_reasoning_accuracy",
            "causal_reasoning_accuracy",
            "emotional_reasoning_accuracy",
            "temporal_reasoning_accuracy",
            "social_reasoning_accuracy",
            "spatial_reasoning_accuracy",
        ]
        for metric_name in metric_names:
            row = metric_table.add_row().cells
            row[0].text = metric_name
            row[1].text = f"{metrics[metric_name]:.4f}"

        document.add_heading("Failure Distribution", level=1)
        if os.path.exists(chart_paths["failure_distribution"]):
            document.add_picture(chart_paths["failure_distribution"], width=Inches(6.2))

        document.add_heading("Category-wise Accuracy", level=1)
        if os.path.exists(chart_paths["category_accuracy"]):
            document.add_picture(chart_paths["category_accuracy"], width=Inches(6.2))

        document.add_heading("Qualitative Failure Analysis", level=1)
        weakness_lines = []
        for failure_type, count in sorted(metrics["failure_distribution"].items(), key=lambda item: item[1], reverse=True)[:5]:
            weakness_lines.append(f"{failure_type}: {count} cases")
        document.add_paragraph("Model weakness analysis: " + ("; ".join(weakness_lines) if weakness_lines else "No failure cases recorded."))

        document.add_heading("Reasoning-type Comparison", level=1)
        if os.path.exists(chart_paths["reasoning_comparison"]):
            document.add_picture(chart_paths["reasoning_comparison"], width=Inches(6.2))

        document.add_heading("Model Performance Comparison", level=1)
        if os.path.exists(chart_paths["model_comparison"]):
            document.add_picture(chart_paths["model_comparison"], width=Inches(6.2))

        document.add_heading("Example Failure Cases", level=1)
        for asset in self._select_report_assets(failure_assets, limit=6):
            document.add_paragraph(asset["record"]["question"])
            document.add_picture(asset["academic_panel"], width=Inches(6.4))

        document.add_heading("Conclusion", level=1)
        document.add_paragraph(
            f"Overall accuracy reached {metrics['overall_accuracy']:.4f}, while the adversarial failure rate was {metrics['adversarial_failure_rate']:.4f}. The most frequent observed failure mode was {self._top_failure(metrics)}."
        )
        try:
            document.save(self.report_path)
        except PermissionError:
            report_dir = os.path.dirname(self.report_path)
            report_name, report_ext = os.path.splitext(os.path.basename(self.report_path))
            fallback_name = f"{report_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{report_ext}"
            document.save(os.path.join(report_dir, fallback_name))

    def _select_report_assets(self, failure_assets, limit=6):
        if limit <= 0:
            return []

        selected = []
        seen_images = set()
        seen_reasoning_types = set()

        for asset in failure_assets:
            record = asset.get("record", {})
            image = record.get("image")
            reasoning_type = record.get("reasoning_type")
            if image in seen_images:
                continue
            if reasoning_type in seen_reasoning_types and len(seen_reasoning_types) < limit:
                continue
            selected.append(asset)
            seen_images.add(image)
            if reasoning_type:
                seen_reasoning_types.add(reasoning_type)
            if len(selected) >= limit:
                return selected

        for asset in failure_assets:
            if asset in selected:
                continue
            record = asset.get("record", {})
            image = record.get("image")
            if image in seen_images:
                continue
            selected.append(asset)
            seen_images.add(image)
            if len(selected) >= limit:
                return selected

        for asset in failure_assets:
            if asset in selected:
                continue
            selected.append(asset)
            if len(selected) >= limit:
                return selected

        return selected

    def _map_reasoning_type(self, port_type):
        port_type = _normalize_text(port_type)
        mapping = {
            "comp": "compositional",
            "compositional": "compositional",
            "causal": "causal",
            "emotion": "emotional",
            "emotional": "emotional",
            "future": "temporal",
            "temporal": "temporal",
            "social": "social",
            "interaction": "social",
            "spatial": "spatial",
        }
        return mapping.get(port_type, port_type or "compositional")

    def _infer_failure_type(self, *, correct, reasoning_type, prediction, ground_truth, source_record):
        if correct:
            return "none"

        prediction_norm = _normalize_text(prediction)
        target_norm = _normalize_text(ground_truth)
        alt_norm = _normalize_text(source_record.get("alt", ""))
        if not prediction_norm or prediction_norm in {"none", "unknown", "unclear", "n a"}:
            return "compositional collapse" if reasoning_type == "compositional" else f"{reasoning_type} failure"
        if _token_overlap(prediction_norm, alt_norm) > max(_token_overlap(prediction_norm, target_norm), 0.0):
            return "distractor confusion"
        if _token_overlap(prediction_norm, target_norm) < 0.2 and len(prediction_norm.split()) >= 3:
            return "hallucination"
        mapping = {
            "compositional": "compositional collapse",
            "causal": "causal reasoning failure",
            "emotional": "emotional misunderstanding",
            "temporal": "temporal confusion",
            "social": "interaction misunderstanding",
            "spatial": "spatial confusion",
        }
        return mapping.get(reasoning_type, "adversarial failure")

    def _failure_summary(self, record):
        if record["failure_type"] == "hallucination":
            return "The prediction diverges from the grounded answer and introduces unsupported content."
        if record["failure_type"] == "distractor confusion":
            return "The model preferred the distractor interpretation over the grounded multi-hop answer."
        return f"The model missed the {record['reasoning_type']} signal needed for the correct multi-hop answer."

    def _bucket_accuracy(self, records):
        return self._safe_rate(sum(1 for record in records if record["correct"]), len(records))

    def _safe_rate(self, numerator, denominator):
        if not denominator:
            return 0.0
        return float(numerator) / float(denominator)

    def _top_failure(self, metrics):
        if not metrics["failure_distribution"]:
            return "no failure mode"
        return max(metrics["failure_distribution"].items(), key=lambda item: item[1])[0]